# coding:utf-8
import os
from combinedocxs import conbine
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION


def execute(mbfile, fpfile, mdfile, resultfile):
    ccppcs11_path = os.path.join(os.getcwd(), "mark_to_docx", "cs11.docx")
    ccppcs3_path = os.path.join(os.getcwd(), "mark_to_docx", "cs3.docx")
    '''封皮文件需要带全部样式'''
    # thread = threading.current_thread()
    # 将markdown文件参考mb.docx模板转化为docx文件
    os.system('pandoc --reference-doc ' + mbfile + ' -o ' + ccppcs11_path + ' ' + mdfile + '')
    # --reference-doc source/nnn.docx
    # 合并封皮
    columnsList = conbine.combineDocx(fpfile, ccppcs11_path, ccppcs3_path)
    print(columnsList)
    # '''设置所有表格样式'''
    # 读取文档
    document = docx.Document(ccppcs3_path)
    tablestyle = document.styles['Table Grid']
    # 获得所有表格对象
    ps = document.tables
    j = 0
    for tb in ps:
        tb.autofit = True
        tb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tb.style = tablestyle
        tb.table_direction = WD_TABLE_DIRECTION.LTR 
        print("%d", len(tb.rows))
        for i in range(columnsList[j]):
            tb.cell(0, i).paragraphs[0].style = 'table'
        j += 1
    # print(columnsList)
    # tb = ps[len(ps) - 1]
    # print("%d", len(tb.rows))
    # for i in range(6):
    #     tb.cell(len(tb.rows) - 1, i).text = ''
    document.save(resultfile)
    os.remove(ccppcs11_path)
    os.remove(ccppcs3_path)


def main():
    os.system('pandoc --reference-doc source/ccppmb.docx -o outfile/testccpp.docx  inputfile/ccpp.md')
    # execute("source/ccppmb.docx", "source/ccppfp.docx", "inputfile/ccpp.md", "outfile/testccpp.docx")


if __name__ == '__main__':
    main()
