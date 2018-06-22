# coding:utf-8
import docx
from zip_De import zip_de
from rename_copy_image import coverFiles
from document_xml import read_xml, find_nodes, write_xml
from xml.etree.ElementTree import Element
import zipfile
import shutil
import os
from platform import system


def combineDocx(filePath1="source/cs1.docx", filePath2="source/cs1.docx", filePath3="source/cs1.docx"):

    # 将源文件复制一份，后面全部操作复制文件
    docx.Document(filePath2).save('cs2catch.docx')
    docx.Document(filePath1).save('cs1catch.docx')

    # 重命名
    os.rename("cs1catch.docx", "test1.zip")
    os.rename("cs2catch.docx", "test2.zip")

    # 获得当前目录
    filedir = os.getcwd()

    # 解压
    zip_de(filedir, 'test1.zip', '/test1/')
    zip_de(filedir, 'test2.zip', '/test2/')

    test2DocumentPath = filedir + "/test2/word/document.xml"
    test2DocumentXmlPath = filedir + "/test2/word/_rels/document.xml.rels"
    test2Content_TypesPath = filedir + "/test2/[Content_Types].xml"
    test2MediaPath = filedir + "/test2/word/media/"
    test1DocumentXmlTreePath = filedir + "/test1/word/_rels/document.xml.rels"
    test1MediaPath = filedir + "/test1/word/media/"

    # 获得cs2中的document.xml对象
    test2DocumentTree = read_xml(test2DocumentPath)
    test2DocumentXmlTree = read_xml(test2DocumentXmlPath)
    test2DocumentXmlNodes = test2DocumentXmlTree.getroot()
    test2Content_TypesTree = read_xml(test2Content_TypesPath)
    test2Content_TypesNodes = test2Content_TypesTree.getroot()
    test1DocumentXmlTree = read_xml(test1DocumentXmlTreePath)
    test1DocumentNodes = test1DocumentXmlTree.getroot()

    w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    wp = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    a = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    pic = "{http://schemas.openxmlformats.org/drawingml/2006/picture}"
    r = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    # 命名空间 和 XPath 搭配使用 报错
    # 2. 属性修改
    # A. 找到父节点
    # test2DocumentTableNodes = find_nodes(test2DocumentTree, w + "body"
    #                 + '/' + w + 'tbl'
    #                 + '/' + w + 'tr'
    #                 + '/' + w + 'trPr'
    #                 + '/' + w + 'cnfStyle'
    #                 )
    # key = w + 'firstRow'

    # 构造一个节点 <w:tblGrid><w:tblGrid/>
    # tblGridnode = Element(w + 'tblGrid')
    
    # 追加表格节点
    columnsList = []
    test2DocumentTableTblPrNodes = find_nodes(test2DocumentTree, w + "body"
                    + '/' + w + 'tbl')
    for tbleseachtree in test2DocumentTableTblPrNodes:
        # test2DocumentTableTrNodes = find_nodes(tbleseachtree, w + 'tr')
        # test2DocumentTableTrPrNodes = find_nodes(tbleseachtree, w + 'tr' + '/' + w + 'trPr')
        columnsList.append(len(find_nodes(tbleseachtree, w + 'tr' + '/' + w + 'tc')))
        # for i in range(1, len(test2DocumentTableTrNodes)):
        #     test2DocumentTableTrNodes[i].append(test2DocumentTableTrPrNodes[0])
        # 修改值
        # test2DocumentTableCnfStyleNodes = find_nodes(tbleseachtree, w + 'tr'
        #             + '/' + w + 'trPr'
        #             + '/' + w + 'cnfStyle')
        # key = w + 'firstRow'
        # for node in test2DocumentTableCnfStyleNodes:
        #     del node 
        # test2DocumentBodyPrNodes.remove(tbleseachtree)
        # tbleseachtree.append(tblGridnode)

    # for node in test2DocumentTableTblPrNodes:
    #     # del node[0]
    #     del node[1]
    #     # del node[2]
    #     # del node[3]
    #     # print(len(node))
    #     # for i in range(len(node) - 3):
    #     #     print(i)
    #     #     del node[i]

    '''追加表格节点结束'''    

    test2DocumentNodes = find_nodes(test2DocumentTree, w + "body"
                    + '/' + w + 'p'
                    + '/' + w + 'r'
                    + '/' + w + 'drawing'
                    + '/' + wp + 'inline'
                    + '/' + a + 'graphic'
                    + '/' + a + 'graphicData'
                    + '/' + pic + 'pic'
                    + '/' + pic + 'blipFill'
                    + '/' + a + 'blip')
    key = r + 'embed'
    i = 1
    # 维护一个字典存放修改前后的embed值
    embeddir = {}
    for node in test2DocumentNodes:
        # 字典的键值不允许重复
        embeddir[node.get(key)] = 'rId100' + str(i)
        i += 1
    print(len(test2DocumentNodes))
    for node in test2DocumentNodes:
        # 获取节点id
        # print("%s =?= %s", node.get(key))
        oldembed = node.get(key)
        del node.attrib[key]
        node.set(key, embeddir[oldembed])
        # print("%s =?= %s", node.get(key))
        # 根据oldembed在document.xml.rels中查找对象
        for resnode in test2DocumentXmlNodes:
            if resnode.get('Id') == oldembed:
                # 重命名test2中的照片名称
                filename_tmp = filedir + "/test2/word/" + resnode.get('Target')
                if os.path.exists(filename_tmp):
                    filenewname = filedir + "/test2/word/media/image" + embeddir[oldembed] + ".png"
                    # 重命名文件丢失
                    os.rename(filename_tmp, filenewname)
                # 修改test2中document.xml.rels文件的节点
                resnode.set('Id', embeddir[oldembed])
                resnode.set('Target', "media/image" + embeddir[oldembed] + ".png")
                # 将此节点追加到test1的document.xml.rels后面
                test1DocumentNodes.append(resnode)
                xmlns = "{http://schemas.openxmlformats.org/package/2006/content-types}"
                test2Content_TypesNewnode = Element(xmlns + 'Override')
                test2Content_TypesNewnode.set('ContentType', 'image/png')
                test2Content_TypesNewnode.set('PartName', '/word/media/image' + embeddir[oldembed] + '.png')
                test2Content_TypesNodes.append(test2Content_TypesNewnode)
        i += 1
    coverFiles(test2MediaPath, test1MediaPath)
    write_xml(test2DocumentXmlTree, test2DocumentXmlPath)
    write_xml(test2DocumentTree, test2DocumentPath)
    write_xml(test2Content_TypesTree, test2Content_TypesPath)
    write_xml(test1DocumentXmlTree, test1DocumentXmlTreePath)

    # 压缩
    myys("test2", "cs2catch")
    myys("test1", "cs1catch")

    # ''''合并文档'''
    filename1 = docx.Document("cs1catch.docx")
    filename2 = docx.Document("cs2catch.docx")
    for elem in filename2.element.body:
        filename1.element.body.append(elem)
    filename1.save(filePath3)
    os.remove("cs1catch.docx")
    os.remove("cs2catch.docx")
    return columnsList


def myys(var1, var2):
    osplat = system()
    filedir = os.getcwd()
    try:
        compression = zipfile.ZIP_DEFLATED
    except:
        compression = zipfile.ZIP_STORED
    FinalZipPath = filedir + '/' + var1 + '/'   # 要进行压缩的文档目录
    if (osplat == 'Windows'):
        FinalZipPath = filedir + '\\' + var1 + '\\'   # 要进行压缩的文档目录
    start = FinalZipPath.rfind(os.sep) + 1
    filename = '' + var1 + '.zip'  # 压缩后的文件名
    z = zipfile.ZipFile(filename, mode="w", compression=compression)
    try:
        for dirpath, dirs, files in os.walk(FinalZipPath):
            for file in files:
                if file == filename or file == "zip.py":
                    continue
                # print(file)
                z_path = os.path.join(dirpath, file)
                z.write(z_path, z_path[start:])
        z.close()
    except:
        if z:
            z.close()
    if (osplat == 'Linux'):
        os.rename(filedir + '/' + var1 + '.zip', "" + var2 + ".docx")
        shutil.rmtree(filedir + '/' + var1 + '')
    if (osplat == 'Windows'):
        os.rename(filedir + '\\' + var1 + '.zip', "" + var2 + ".docx")
        shutil.rmtree(filedir + '\\' + var1 + '')
