# coding:utf-8
import docx
import sys
import os
import re


def createmd():
    ''''生成md出版文件'''

    document = docx.Document(ur"C:/Users/ren.bo2/Desktop/CCPP测试方案新版.docx")
    paragraphs = document.paragraphs
    ps = document.tables
    maxi = len(ps)
    print(maxi)
    # # 遍历每段，在每段中执行替换动作
    f = open('a.md', 'w')
    old = sys.stdout
    # 将当前系统输出储存到一个临时变量中
    sys.stdout = f
    dirll = {'Normal': '##### ', 'Heading 1': '# ', 'Heading 2': '## ', 'Heading 3': '### '}
    i = 0
    for para in paragraphs:
        j = 0
        if len(para.text) != 0:
            # 判断是否包含两个*@@@@@@@表格
            ddd = para.text.replace('*', '@@@', len(para.text))
            print(dirll[para.style.name] + ddd.encode('utf-8'))
            if ddd.find(u"@@@@@@@表格") != -1 and i < maxi:
                print
                for row in ps[i].rows:
                    hang = '| '
                    for cell in row.cells:
                        hang = hang + cell.text + ' |'
                    print(hang.encode('utf-8'))
                    if j == 0:
                        hang = '|'
                        for cell in row.cells:
                            hang = hang + ':------|'
                        print(hang.encode('utf-8'))
                        j += 1
                i += 1
                print
    sys.stdout = old
    f.close()
    # 遍历每段，在每段中执行替换动作
    f = open('result.md', 'w')
    old = sys.stdout
    # 将当前系统输出储存到一个临时变量中
    sys.stdout = f
    for line in open("a.md"):
        ddd = line.replace('@@@@@@@表格', '', 1)
        ddd = ddd.replace('@', '', len(line))
        print ddd,
    sys.stdout = old
    # 还原原系统输出
    f.close()
    os.remove('a.md')


''''解析md文件替换全部需要的值'''

plandocxcolumdir = {'company.company_name': 1,
                    u'ccpp_questionnaire.engine_power': 2,
                    'ccpp_questionnaire.recent_steam_flow_range_1': 3,
                    'ccpp_questionnaire.steam_pressure_level_1': 4,
                    'ccpp_questionnaire.steam_temperature_level_1': 5,
                    'ccpp_questionnaire.methane_design': 6,
                    'ccpp_questionnaire.ethane_design': 7,
                    'ccpp_questionnaire.ethylene_design': 8,
                    'ccpp_questionnaire.propylene_design': 9,
                    'ccpp_questionnaire.propane_design': 10,
                    'ccpp_questionnaire.butene_design': 11,
                    'ccpp_questionnaire.i_isobutane_design': 12,
                    'ccpp_questionnaire.n_isobutane_design': 13,
                    'ccpp_questionnaire.pentane_design': 14,
                    'ccpp_questionnaire.carbon6_design': 15,
                    'ccpp_questionnaire.hydrogen_design': 16,
                    'ccpp_questionnaire.helium_design': 17,
                    'ccpp_questionnaire.nitrogen_design': 18,
                    'ccpp_questionnaire.carbon_monoxide_design': 19,
                    'ccpp_questionnaire.carbon_dioxide_design': 20,
                    'ccpp_questionnaire.hydrogen_sulfide_design': 21,
                    'ccpp_questionnaire.oxygen_design': 22,
                    'ccpp_questionnaire.water_design': 23,
                    'ccpp_questionnaire.high_calorific_value_design': 15,
                    'ccpp_questionnaire.price_design': 45,
                    'ccpp_questionnaire.local_avg_hight': 12,
                    'ccpp_questionnaire.year_avg_temperate': 45,
                    'ccpp_questionnaire.summer_avg_temperate': 13,
                    'ccpp_questionnaire.winter_avg_temperate': 45,
                    'ccpp_questionnaire.year_avg_press': 78,
                    'ccpp_questionnaire.summer_avg_press': 20,
                    'ccpp_questionnaire.winter_avg_press': 11,
                    'ccpp_questionnaire.year_avg_humidity': 78,
                    'ccpp_questionnaire.electric_load_demand': 78,

                    'ccpp_ccpp.engine_model': 55,
                    'ccpp_ccpp.high_terminal_temperature_difference': 55,
                    'ccpp_ccpp.low_superheater_effluent_smoke_enthalpy': 55,
                    'ccpp_ccpp.engine_power': 55,
                    'ccpp_ccpp.high_steam_enthalpy': 55,
                    'ccpp_ccpp.engine_num': 55,
                    'ccpp_ccpp.engine_heat_consumption_rate': 55,
                    'ccpp_ccpp.engine_efficiency': 55,
                    'ccpp_ccpp.engine_exhuast_gas_temperature': 55,
                    'ccpp_ccpp.individual_gas_consumption': 66,
                    'ccpp_ccpp.engine_exhuast_gas_flux': 66,
                    'ccpp_ccpp.sp_steam_pressure': 55,
                    'ccpp_ccpp.sp_steam_temperature': 55,
                    'ccpp_ccpp.sp_low_gas_production': 55,
                    'ccpp_ccpp.sp_low_feedwater_temperature': 55,
                    'ccpp_ccpp.high_steam_pressure': 55,
                    'ccpp_ccpp.high_steam_temperature': 55,
                    'ccpp_ccpp.high_gas_production': 55,
                    'ccpp_ccpp.high_economizer_effluent_water_temperature': 55,
                    'ccpp_ccpp.low_drum_pressure': 55,
                    'ccpp_ccpp.low_effluent_smoke_temperature': 55,
                    'ccpp_ccpp.low_gas_production': 55,
                    'ccpp_ccpp.low_feedwater_temperature': 55,
                    'ccpp_ccpp.boiler_single_or_dula_pressure': 'singlepot'

                    }
plandocxcolumdir['individual_gas_consumption_sum'] = plandocxcolumdir['ccpp_ccpp.individual_gas_consumption'] * plandocxcolumdir['ccpp_ccpp.engine_num']
plandocxcolumdir['engine_exhuast_gas_flux_sum'] = plandocxcolumdir['ccpp_ccpp.engine_exhuast_gas_flux'] * plandocxcolumdir['ccpp_ccpp.engine_num']
plandocxcolumdir['electric_load_demand0.8'] = plandocxcolumdir['ccpp_questionnaire.electric_load_demand'] * 0.8
plandocxcolumdir['engine_power0.8'] = plandocxcolumdir['ccpp_ccpp.engine_power'] * 0.8
plandocxcolumdir['individual_gas_consumption8000'] = plandocxcolumdir['ccpp_ccpp.individual_gas_consumption'] * 0.8
plandocxcolumdir['individual_gas_consumption8000engine_power0.8'] = plandocxcolumdir['ccpp_ccpp.individual_gas_consumption'] * 8000 / plandocxcolumdir['ccpp_ccpp.engine_power'] * 0.8
if 'singlepot' == plandocxcolumdir['ccpp_ccpp.boiler_single_or_dula_pressure']:
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1'] = plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1锅炉额定蒸发量'] = plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1 * 1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1厂内汽水循环正常损失'] = plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1 * 0.03
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1锅炉正常排污损失'] = plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1 * 0.02
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1水处理系统耗水量'] = round(plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1 * 0.03, 0)
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1其它不可预计用水损失'] = round(plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1 * 0.03, 0)
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1启动或事故增加用水量'] = plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1 * 0.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1外供汽损失'] = plandocxcolumdir['ccpp_ccpp.sp_low_gas_production'] * 1.1 * 0.3
else:
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1锅炉额定蒸发量'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1厂内汽水循环正常损失'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1锅炉正常排污损失'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1水处理系统耗水量'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1其它不可预计用水损失'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1启动或事故增加用水量'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
    plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1外供汽损失'] = (plandocxcolumdir['ccpp_ccpp.high_gas_production'] + plandocxcolumdir['ccpp_ccpp.low_gas_production']) * 1.1
plandocxcolumdir[u'厂内汽水循环正常损失：+锅炉正常排污损失+水处理系统耗水量+其它不可预计用水损失：+外供汽损失t/h =12.8'] = plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1厂内汽水循环正常损失'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1锅炉正常排污损失'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1水处理系统耗水量'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1其它不可预计用水损失'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1外供汽损失']
plandocxcolumdir[u'厂内汽水循环正常损失：+锅炉正常排污损失+水处理系统耗水量+其它不可预计用水损失：+外供汽损失+启动或事故增加用水量t/h =16.4'] = plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1厂内汽水循环正常损失'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1锅炉正常排污损失'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1水处理系统耗水量'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1其它不可预计用水损失'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1外供汽损失'] + plandocxcolumdir[u'(high_gas_production+low_gas_production)1.1andsp_low_gas_production1.1外供汽损失']
plandocxcolumdir[u'(锅炉补给水处理系统正常出力+锅炉补给水处理系统最大出力)/2'] = (plandocxcolumdir[u'厂内汽水循环正常损失：+锅炉正常排污损失+水处理系统耗水量+其它不可预计用水损失：+外供汽损失t/h =12.8'] + plandocxcolumdir[u'厂内汽水循环正常损失：+锅炉正常排污损失+水处理系统耗水量+其它不可预计用水损失：+外供汽损失+启动或事故增加用水量t/h =16.4']) * 0.5
plandocxcolumdir[u'(锅炉补给水处理系统正常出力+锅炉补给水处理系统最大出力)2.5取整'] = (plandocxcolumdir[u'厂内汽水循环正常损失：+锅炉正常排污损失+水处理系统耗水量+其它不可预计用水损失：+外供汽损失t/h =12.8'] + plandocxcolumdir[u'厂内汽水循环正常损失：+锅炉正常排污损失+水处理系统耗水量+其它不可预计用水损失：+外供汽损失+启动或事故增加用水量t/h =16.4']) * 2.5


def mdreplaceold(filepath, y, z):
    with open(filepath, "r",) as f:
        # readlines以列表的形式将文件读出
        lines = f.readlines()
    with open(filepath, "w",) as f_w:
        # 定义一个数字，用来记录在读取文件时在列表中的位置
        for line in lines:
            if y in line.decode('utf8'):
                line = line.decode('utf8').replace(unicode(y), z)
                f_w.write(line.encode('utf8'))
            else:
                f_w.write(line)


def mdreplace(lines, y, z):
    newlines = ""
    for line in lines:
        if y in line.decode('utf8'):
            line = line.decode('utf8').replace(unicode(y), z)
        newlines += line
    return newlines
                

def main():
    createmd()
    # for key in plandocxcolumdir:
    #     mdreplace('result.md', key, str(plandocxcolumdir[key]))
    

if __name__ == '__main__':
    main()
